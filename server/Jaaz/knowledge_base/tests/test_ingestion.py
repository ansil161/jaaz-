"""Detection, extraction, normalisation and chunking.

These are the parts with no database and no network, so they are tested
directly rather than through the API.
"""

import io
from unittest import TestCase

from knowledge_base.config import ChunkingConfig
from knowledge_base.ingestion.chunker import TextChunker
from knowledge_base.ingestion.detection import (
    DOCX,
    MARKDOWN,
    PDF,
    PLAIN_TEXT,
    UnsupportedDocumentError,
    detect,
)
from knowledge_base.ingestion.extractors.base import (
    ExtractedSegment,
    ExtractionError,
)
from knowledge_base.ingestion.extractors.docx import DocxExtractor
from knowledge_base.ingestion.extractors.pdf import PdfExtractor
from knowledge_base.ingestion.extractors.plain_text import (
    MarkdownExtractor,
    PlainTextExtractor,
)
from knowledge_base.ingestion.normalizer import normalize

from .support import docx_bytes, pdf_bytes


def stream(payload):
    return io.BytesIO(payload)


class DetectionTests(TestCase):
    def test_a_pdf_is_detected_from_its_magic_bytes(self):
        self.assertIs(detect(stream(pdf_bytes(['hello'])), 'manual.pdf'), PDF)

    def test_a_docx_is_detected_by_its_package_entry(self):
        self.assertIs(detect(stream(docx_bytes(['hello'])), 'policy.docx'), DOCX)

    def test_text_is_detected(self):
        self.assertIs(detect(stream(b'plain words'), 'notes.txt'), PLAIN_TEXT)

    def test_markdown_is_told_from_text_by_its_extension(self):
        # They are byte-identical; the extension is the only signal, and it
        # is the one thing the extension is trusted for.
        self.assertIs(detect(stream(b'# Title'), 'guide.md'), MARKDOWN)

    def test_the_browser_content_type_is_never_consulted(self):
        """A PDF announced as text is still a PDF, and vice versa.

        `detect` is not even given the client's Content-Type — there is no
        parameter for it.
        """
        self.assertIs(detect(stream(pdf_bytes(['x'])), 'thing.pdf'), PDF)

    def test_a_docx_renamed_as_a_pdf_is_rejected(self):
        with self.assertRaises(UnsupportedDocumentError) as caught:
            detect(stream(docx_bytes(['hello'])), 'invoice.pdf')
        self.assertIn('DOCX', caught.exception.message)

    def test_a_pdf_renamed_as_a_docx_is_rejected(self):
        with self.assertRaises(UnsupportedDocumentError):
            detect(stream(pdf_bytes(['hello'])), 'invoice.docx')

    def test_binary_content_is_rejected(self):
        with self.assertRaises(UnsupportedDocumentError):
            detect(stream(b'\x7fELF\x02\x01\x01\x00rest'), 'payload.txt')

    def test_a_zip_that_is_not_a_docx_is_rejected(self):
        import zipfile

        buffer = io.BytesIO()
        with zipfile.ZipFile(buffer, 'w') as archive:
            archive.writestr('hello.txt', 'not a word document')
        with self.assertRaises(UnsupportedDocumentError):
            detect(stream(buffer.getvalue()), 'archive.docx')

    def test_an_executable_extension_over_text_is_treated_as_text(self):
        """Documented behaviour, not an oversight.

        The bytes are text, and text is what gets ingested. The upload is
        re-keyed with a .txt extension in storage, is never served, and is
        never executed — so the extension carries no authority either way.
        """
        self.assertIs(detect(stream(b'echo hello'), 'script.exe'), PLAIN_TEXT)

    def test_a_null_byte_disqualifies_text(self):
        with self.assertRaises(UnsupportedDocumentError):
            detect(stream(b'mostly text\x00but not'), 'notes.txt')


class ExtractorTests(TestCase):
    def test_pdf_extraction_records_the_page_number(self):
        extracted = PdfExtractor().extract(
            stream(pdf_bytes(['First page text', 'Second page text']))
        )
        self.assertEqual(len(extracted.segments), 2)
        self.assertEqual(extracted.segments[0].metadata['page'], 1)
        self.assertEqual(extracted.segments[1].metadata['page'], 2)
        self.assertIn('Second page', extracted.text)

    def test_a_pdf_with_no_text_layer_is_an_error_not_an_empty_document(self):
        # A scan. Silently ingesting it would produce a document that is
        # "Ready" and retrieves nothing.
        with self.assertRaises(ExtractionError) as caught:
            PdfExtractor().extract(stream(pdf_bytes([''])))
        self.assertIn('OCR', caught.exception.message)

    def test_a_corrupt_pdf_reports_a_safe_message(self):
        with self.assertRaises(ExtractionError) as caught:
            PdfExtractor().extract(stream(b'%PDF-1.4\nnot really a pdf'))
        self.assertIn('could not be read', caught.exception.message)

    def test_docx_extraction_reads_paragraphs(self):
        extracted = DocxExtractor().extract(
            stream(docx_bytes(['Returns within 14 days.', 'Warranty is 3 years.']))
        )
        self.assertIn('Returns within 14 days.', extracted.text)
        self.assertIn('Warranty is 3 years.', extracted.text)

    def test_docx_extraction_includes_tables(self):
        from docx import Document as DocxDocument

        document = DocxDocument()
        document.add_paragraph('Specifications')
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = 'Model'
        table.cell(0, 1).text = 'X200'
        table.cell(1, 0).text = 'Warranty'
        table.cell(1, 1).text = '36 months'
        buffer = io.BytesIO()
        document.save(buffer)

        extracted = DocxExtractor().extract(stream(buffer.getvalue()))

        # The row has to stay on one line, or "Warranty" and "36 months" end
        # up in different chunks and neither answers the question.
        self.assertIn('Warranty | 36 months', extracted.text)

    def test_markdown_syntax_is_stripped_but_words_are_kept(self):
        source = (
            '# Heading\n\n'
            'See the [product guide](https://example.com/guide) for detail.\n\n'
            '![a diagram](https://example.com/img.png)\n'
        )
        text = MarkdownExtractor().extract(stream(source.encode())).text

        self.assertIn('Heading', text)
        self.assertIn('product guide', text)
        self.assertNotIn('https://example.com', text)
        self.assertNotIn('#', text)

    def test_markdown_sections_carry_their_heading(self):
        """Each section becomes a segment tagged with the heading above it.

        This is what lets an answer cite "JAAZ Warranty, Exclusions" rather
        than just naming the document. The extractor used to return the whole
        file as one untagged segment, so every chunk reached the vector store
        with no heading and a reader checking a citation had to search the
        file by hand.
        """
        source = (
            'Intro text with no heading.\n\n'
            '# Warranty\n\n'
            'Thirty-six months from handover.\n\n'
            '## Exclusions\n\n'
            'Water damage is not covered.\n'
        )

        segments = MarkdownExtractor().extract(stream(source.encode())).segments
        headings = [segment.metadata.get('heading') for segment in segments]

        self.assertEqual(headings, [None, 'Warranty', 'Exclusions'])
        # The heading stays in the body too: it is usually the best summary
        # of the passage beneath it, and it helps the chunk retrieve.
        self.assertIn('Warranty', segments[1].text)
        self.assertIn('Thirty-six months', segments[1].text)

    def test_markdown_without_headings_is_still_one_segment(self):
        segments = MarkdownExtractor().extract(
            stream(b'Just prose. No headings anywhere.')
        ).segments

        self.assertEqual(len(segments), 1)
        self.assertEqual(segments[0].metadata, {})

    def test_an_empty_text_file_is_an_error(self):
        with self.assertRaises(ExtractionError):
            PlainTextExtractor().extract(stream(b'   \n  \n'))


class NormalizerTests(TestCase):
    def test_soft_hyphens_and_zero_width_characters_are_removed(self):
        self.assertEqual(normalize('war\xadran​ty'), 'warranty')

    def test_a_word_split_across_a_line_break_is_rejoined(self):
        self.assertEqual(normalize('hyphen-\nated word'), 'hyphenated word')

    def test_a_genuine_hyphen_before_a_capital_is_left_alone(self):
        self.assertEqual(normalize('Anglo-\nSaxon'), 'Anglo- Saxon')

    def test_single_newlines_become_spaces_and_double_ones_survive(self):
        self.assertEqual(
            normalize('one\ntwo\n\nthree'),
            'one two\n\nthree',
        )

    def test_runs_of_blank_lines_collapse(self):
        self.assertEqual(normalize('a\n\n\n\n\nb'), 'a\n\nb')

    def test_ligatures_are_folded(self):
        self.assertEqual(normalize('the ﬁnal oﬃce'), 'the final office')

    def test_empty_input_is_empty_output(self):
        self.assertEqual(normalize(''), '')


class ChunkerTests(TestCase):
    def config(self, **overrides):
        values = {'size': 200, 'overlap': 40, 'minimum': 30, 'maximum': 400}
        values.update(overrides)
        return ChunkingConfig(**values)

    def chunk(self, text, **overrides):
        chunker = TextChunker(self.config(**overrides))
        return chunker.chunk([ExtractedSegment(text=text, metadata={})])

    def test_a_chunk_spanning_sections_is_labelled_with_the_first(self):
        """The heading has to name where the passage begins.

        REGRESSION. The merge used to keep the last heading it saw, so a
        chunk that opened with a specification table and ran on into the next
        section was cited as that next section — sending a reader checking
        the citation to a place the table is not.
        """
        chunker = TextChunker(self.config(size=400, maximum=800))
        chunks = chunker.chunk([
            ExtractedSegment(text='Typical specification. Screen is 2.35:1.',
                             metadata={'heading': 'Typical specification'}),
            ExtractedSegment(text='Look at another solution if the room stays a '
                                  'living room.',
                             metadata={'heading': 'Look at another solution if'}),
        ])

        self.assertEqual(chunks[0].metadata['heading'], 'Typical specification')

    def test_a_chunk_inherits_a_heading_from_a_later_piece_when_it_starts_bare(self):
        """A chunk beginning with unlabelled text still gets the next heading.

        Better than no heading at all: it is the section the body of the
        chunk actually sits in.
        """
        chunker = TextChunker(self.config(size=400, maximum=800))
        chunks = chunker.chunk([
            ExtractedSegment(text='Preamble with no heading.', metadata={}),
            ExtractedSegment(text='Warranty runs thirty-six months.',
                             metadata={'heading': 'Warranty'}),
        ])

        self.assertEqual(chunks[0].metadata['heading'], 'Warranty')

    def test_short_text_is_a_single_chunk(self):
        chunks = self.chunk('A short paragraph.')
        self.assertEqual(len(chunks), 1)
        self.assertEqual(chunks[0].index, 0)
        self.assertEqual(chunks[0].content, 'A short paragraph.')

    def test_paragraphs_are_packed_up_to_the_target_size(self):
        paragraph = 'word ' * 20  # ~100 characters
        chunks = self.chunk('\n\n'.join([paragraph] * 6))

        self.assertGreater(len(chunks), 1)
        self.assertTrue(all(chunk.content for chunk in chunks))
        self.assertEqual([chunk.index for chunk in chunks], list(range(len(chunks))))

    def test_later_chunks_carry_overlap_from_the_previous_one(self):
        first = 'alpha ' * 30
        second = 'beta ' * 30
        chunks = self.chunk(f'{first.strip()}\n\n{second.strip()}')

        self.assertGreater(len(chunks), 1)
        # The tail of chunk 0 reappears at the head of chunk 1, which is what
        # keeps a fact spanning the boundary retrievable from either side.
        self.assertIn('alpha', chunks[1].content)

    def test_a_paragraph_over_the_maximum_is_split_on_sentences(self):
        sentence = 'This is a sentence about warranties. '
        chunks = self.chunk(sentence * 40, maximum=300)

        self.assertTrue(all(len(chunk.content) <= 700 for chunk in chunks))
        self.assertGreater(len(chunks), 3)

    def test_text_with_no_spaces_still_terminates_and_is_bounded(self):
        # The pathological input: no paragraph, sentence or word boundary
        # anywhere. The hard split has to advance regardless.
        chunks = self.chunk('x' * 5000, maximum=400)

        self.assertGreater(len(chunks), 5)
        self.assertTrue(all(chunk.content for chunk in chunks))

    def test_a_tiny_trailing_chunk_is_folded_into_the_one_before_it(self):
        body = 'word ' * 60
        chunks = self.chunk(f'{body.strip()}\n\nPage 3 of 3')

        self.assertNotEqual(chunks[-1].content.strip(), 'Page 3 of 3')
        self.assertIn('Page 3 of 3', chunks[-1].content)

    def test_page_metadata_from_several_segments_is_merged(self):
        chunker = TextChunker(self.config(size=1000, maximum=2000))
        chunks = chunker.chunk(
            [
                ExtractedSegment(text='First page body.', metadata={'page': 1}),
                ExtractedSegment(text='Second page body.', metadata={'page': 2}),
            ]
        )
        self.assertEqual(chunks[0].metadata['pages'], [1, 2])

    def test_token_count_is_estimated_for_every_chunk(self):
        chunks = self.chunk('word ' * 100)
        self.assertTrue(all(chunk.token_count > 0 for chunk in chunks))

    def test_a_bad_configuration_is_rejected_when_it_is_built(self):
        with self.assertRaises(ValueError):
            ChunkingConfig(size=100, overlap=100, minimum=10, maximum=200)
        with self.assertRaises(ValueError):
            ChunkingConfig(size=0, overlap=0, minimum=0, maximum=10)
