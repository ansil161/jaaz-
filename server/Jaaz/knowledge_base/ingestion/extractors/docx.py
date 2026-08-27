"""DOCX extraction.

Tables are included, not skipped. In the kind of document a knowledge base
holds — policies, specifications, price lists — the table is frequently where
the answer is, and a parser that walks only `document.paragraphs` silently
drops it.

Rows are flattened to ` | `-separated cells. That keeps a row's values
associated with each other in the embedding, which is what a question like
"what is the warranty on the X200" needs.
"""

from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError

from .base import ExtractedDocument, ExtractedSegment, ExtractionError


class DocxExtractor:
    kind_key = 'docx'

    def extract(self, stream) -> ExtractedDocument:
        stream.seek(0)

        try:
            document = DocxDocument(stream)
        except (PackageNotFoundError, KeyError, ValueError, OSError) as exc:
            raise ExtractionError(
                'This DOCX file could not be read. It may be corrupt, or it '
                'may be an older .doc file saved with the wrong extension.'
            ) from exc

        segments = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            metadata = {}
            style = getattr(paragraph.style, 'name', '') or ''
            if style.startswith('Heading'):
                # Recorded so a chunk can be labelled with the section it
                # came from once the chunker groups them.
                metadata['heading'] = text
            segments.append(ExtractedSegment(text=text, metadata=metadata))

        for index, table in enumerate(document.tables, start=1):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    rows.append(' | '.join(cells))
            if rows:
                segments.append(
                    ExtractedSegment(
                        text='\n'.join(rows), metadata={'table': index}
                    )
                )

        if not segments:
            raise ExtractionError('No text could be read from this document.')

        return ExtractedDocument(segments)
